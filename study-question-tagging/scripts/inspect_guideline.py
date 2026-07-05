#!/usr/bin/env python3
"""Dump a guideline / manuscript document so you can locate its clinical questions.

Clinical questions live in different places in different guidelines: a "Guideline
Questions" overview table, a per-section recommendations table, or inline under
"Clinical Question N" headings followed by the question text and a recommendation.
This script surfaces all of those so you can read the questions verbatim instead of
guessing — getting a question's exact wording, population, and scope right is what
makes the downstream study assignment defensible.

Usage:
  python inspect_guideline.py <doc.docx> [--outline] [--tables] [--questions]
                                          [--dump-paras OUT.txt]

  (no flags)      print an outline (headings) + a table inventory + question guesses
  --outline       headings only
  --tables        full dump of every table (index, dimensions, cell text)
  --questions     heuristic extraction of likely clinical-question text
  --dump-paras F  write every non-empty paragraph (with index + style) to file F
                  so an agent can Read the relevant span in full

Supports .docx (python-docx). For .pdf guidelines, read the file directly with the
Read tool — this script is for the common .docx manuscript case. .txt is dumped raw.
"""
import argparse, os, re, sys

QUESTION_HINT = re.compile(r"(clinical question|guideline question|^\s*CQ\s*\d|key question|PICO)", re.I)


def load_docx(path):
    try:
        import docx
    except ImportError:
        sys.exit("python-docx not installed: pip install python-docx")
    return docx.Document(path)


def outline(d):
    print("=== OUTLINE (headings) ===")
    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        sty = (p.style.name if p.style else "") or ""
        if sty.lower().startswith(("heading", "title")):
            print(f"  [{i}] ({sty}) {t[:120]}")


def table_inventory(d):
    print("\n=== TABLE INVENTORY ===")
    for ti, t in enumerate(d.tables):
        nr, nc = len(t.rows), len(t.columns)
        try:
            first = t.rows[0].cells[0].text.strip().replace("\n", " ")[:70]
        except Exception:
            first = ""
        print(f"  TABLE {ti}: {nr}x{nc} | first cell: {first}")


def dump_tables(d):
    for ti, t in enumerate(d.tables):
        print(f"\n===== TABLE {ti} ({len(t.rows)}x{len(t.columns)}) =====")
        for r in t.rows:
            cells = [c.text.strip().replace("\n", " / ") for c in r.cells]
            print(" | ".join(cells))
            print("-" * 60)


def guess_questions(d):
    """Surface text that looks like a clinical question.

    Two signals: (1) a paragraph near a 'Clinical Question'/'CQ N' marker, and
    (2) any table whose header row mentions 'Clinical Question' — dump its rows.
    This is a starting point; always confirm against the verbatim document.
    """
    print("\n=== LIKELY CLINICAL QUESTIONS (heuristic — verify against the doc) ===")
    paras = d.paragraphs
    seen = set()
    for i, p in enumerate(paras):
        t = p.text.strip()
        if not t:
            continue
        if QUESTION_HINT.search(t) or (t.endswith("?") and len(t) > 40):
            # print this paragraph and, if it's a 'Clinical Question N' marker,
            # the following non-empty paragraph (which is usually the question text)
            if i in seen:
                continue
            print(f"  [{i}] {t}")
            seen.add(i)
            if re.match(r"^\s*(clinical|guideline)\s+question", t, re.I) or re.match(r"^\s*CQ\s*\d", t, re.I):
                for j in range(i + 1, min(i + 3, len(paras))):
                    nt = paras[j].text.strip()
                    if nt:
                        print(f"      -> [{j}] {nt}")
                        seen.add(j)
                        break
    # tables that look like question lists
    for ti, t in enumerate(d.tables):
        try:
            hdr = " ".join(c.text for c in t.rows[0].cells).lower()
        except Exception:
            hdr = ""
        if "clinical question" in hdr or "guideline question" in hdr:
            print(f"\n  -- TABLE {ti} appears to list questions --")
            for r in t.rows:
                cells = [c.text.strip().replace("\n", " ") for c in r.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    print(f"     {line[:300]}")


def dump_paras(d, out):
    n = 0
    with open(out, "w") as f:
        for i, p in enumerate(d.paragraphs):
            t = p.text.strip()
            if t:
                sty = (p.style.name if p.style else "") or "Normal"
                f.write(f"[{i}|{sty}] {t}\n")
                n += 1
    print(f"wrote {n} non-empty paragraphs to {out}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("doc")
    ap.add_argument("--outline", action="store_true")
    ap.add_argument("--tables", action="store_true")
    ap.add_argument("--questions", action="store_true")
    ap.add_argument("--dump-paras")
    a = ap.parse_args()

    ext = os.path.splitext(a.doc)[1].lower()
    if ext == ".txt":
        with open(a.doc) as f:
            print(f.read())
        return
    if ext != ".docx":
        sys.exit(f"Unsupported guideline type {ext}. For PDFs, read with the Read tool.")

    d = load_docx(a.doc)
    if a.dump_paras:
        dump_paras(d, a.dump_paras)
        return
    if a.tables:
        dump_tables(d)
        return
    if a.questions:
        guess_questions(d)
        return
    if a.outline:
        outline(d)
        return
    # default: a bit of everything
    outline(d)
    table_inventory(d)
    guess_questions(d)


if __name__ == "__main__":
    main()
